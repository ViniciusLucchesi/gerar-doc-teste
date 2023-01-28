import re
from pathlib import Path
from docx import Document
from win32api import GetUserNameEx, GetUserName
from datetime import datetime
from scripts.path import FindFiles


class GerarDocTeste:
    def __init__(self, change:str='padrao'):
        self.change = change.upper()
        self.author = ''
        self.today = datetime.today().strftime('%d/%m/%Y')
        self.doc_file = ''
        self.is_valid = self._validate_change_number()
        if self.is_valid:
            self.doc_number = self._get_doc_number()
        else:
            self.doc_number = ''


    def _validate_change_number(self) -> bool:
        found = re.search('^(CHG[0-9]{7})$', self.change)
        if found == None:
            return False
        return True      


    def _get_current_author_format(self, author_option:str) -> None:
        if author_option == 'Username':
            self.author = GetUserName()
        elif author_option == 'Nome completo':
            self.author = GetUserNameEx(3)


    def _get_doc_number(self) -> str:
        file = FindFiles(self.change)
        file.find_documents()
        file.find_last_document_number()
        return file.doc_number
    

    def _get_file_name(self, file_path) -> str:
        file_name = re.sub('/', '###', file_path).split('###')[-1]
        return file_name


    def create_word_doc(self, doc_name:str, directory:str, author_option:str) -> None:
        doc = Document(f'word_template/{doc_name}')
        self._get_current_author_format(author_option)

        metadata = doc.core_properties
        metadata.author = self.author
        metadata.title = self.change
        header = doc.sections[0].header       

        self.replace_information(header)
        self.replace_information(doc)

        self.doc_file = f'{self.change}_{self.doc_number}.docx'
        doc.save(f'{directory}{self.doc_file}')


    def replace_information(self, section) -> None:
        for paragraph in section.paragraphs:
            for run in paragraph.runs:
                text = run.text
                text = text.replace('TITLE', self.change)
                text = text.replace('AUTHOR', self.author)
                text = text.replace('DATE', self.today)
                run.text = text


    def save_document(self, doc_path) -> None:
        default_path = re.sub('(")|(")', '', doc_path)
        default_path = re.sub('\\\\', '/', default_path)
        doc = Document(default_path)
        file_name = self._get_file_name(default_path)
        full_path = str(Path('word_template', file_name))
        doc.save(full_path)
